import os
import PyPDF2
import docx
import pandas as pd
import pypandoc
from io import BytesIO
from deep_translator import GoogleTranslator
from azure.ai.inference import ChatCompletionsClient
from azure.ai.inference.models import SystemMessage, UserMessage
from azure.core.credentials import AzureKeyCredential

# Initialize Azure client
endpoint = os.getenv("AZURE_INFERENCE_SDK_ENDPOINT", "https://aihubdeepseek5311705507.services.ai.azure.com/models")
model_name = os.getenv("DEPLOYMENT_NAME", "DeepSeek-R1")
key = os.getenv("AZURE_INFERENCE_SDK_KEY", "1K2kkjX1MOoB1BQJfm0LbXXZamYheCQ4qZaO7oAroS9IbVTsJYBkJQQJ99BCACYeBjFXJ3w3AAAAACOGuoUz")
client = ChatCompletionsClient(endpoint=endpoint, credential=AzureKeyCredential(key))

def get_chat_response(user_input):
    response = client.complete(
        messages=[
            SystemMessage(content="You are a helpful assistant."),
            UserMessage(content=user_input)
        ],
        model=model_name,
        max_tokens=1000
    )
    return response.choices[0].message.content

def summarize_file(uploaded_file):
    # Extract text based on file type
    file_type = uploaded_file.type
    text = ""

    if file_type == "application/pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        for page in reader.pages:
            text += page.extract_text()
    elif file_type == "text/plain":
        text = uploaded_file.read().decode("utf-8")
    elif file_type == "application/msword":
        # Convert .doc to .docx using pypandoc
        docx_file = "temp.docx"
        with open(docx_file, "wb") as f:
            f.write(uploaded_file.getbuffer())
        pypandoc.convert_file(docx_file, "docx", outputfile=docx_file)
        doc = docx.Document(docx_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        os.remove(docx_file)  # Clean up the temporary file
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Read .docx file directly
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif file_type == "text/csv":
        df = pd.read_csv(uploaded_file)
        text = df.to_string()

    # Summarize the extracted text using DeepSeek
    summary = get_chat_response(f"Summarize the following text in 100 words or less:\n{text}")

    # Translate the summary to Spanish
    translated_summary = GoogleTranslator(source="auto", target="es").translate(summary)

    # Create a .docx file with the translated summary
    doc = docx.Document()
    doc.add_paragraph(translated_summary)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return translated_summary, doc_buffer